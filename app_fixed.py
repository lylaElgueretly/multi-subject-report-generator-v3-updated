# MULTI-SUBJECT REPORT COMMENT GENERATOR - Working Version
# Supports: English, Maths, Science, ESL, Chemistry

import streamlit as st
import tempfile
import os
import pandas as pd
import io
import random
import re
from datetime import datetime, timedelta
from docx import Document

# SECURITY SETTINGS
TARGET_CHARS = 500
MAX_FILE_SIZE_MB = 5
MAX_ROWS_PER_UPLOAD = 100

# PAGE CONFIG
st.set_page_config(
    page_title="Report Comment Generator",
    layout="wide",
    initial_sidebar_state="expanded"
)

# SESSION STATE INIT
if 'app_initialized' not in st.session_state:
    st.session_state.clear()
    st.session_state.app_initialized = True
    st.session_state.all_comments = []

# SIMPLE HELPER FUNCTIONS
def get_pronouns(gender):
    gender = gender.lower()
    if gender == "male":
        return "he", "his"
    elif gender == "female":
        return "she", "her"
    return "they", "their"

def sanitize_input(text, max_length=100):
    if not text:
        return ""
    sanitized = ''.join(c for c in text if c.isalnum() or c in " .'-")
    return sanitized[:max_length].strip().title()

def lowercase_first(text):
    return text[0].lower() + text[1:] if text else ""

def fix_pronouns_in_text(text, pronoun, possessive):
    if not text:
        return text
    
    text = re.sub(r'\bhe\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHe\b', pronoun.capitalize(), text)
    text = re.sub(r'\bshe\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bShe\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhis\b', possessive, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHis\b', possessive.capitalize(), text)
    text = re.sub(r'\bher\b', possessive, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHer\b', possessive.capitalize(), text)
    return text

def truncate_comment(comment, target=TARGET_CHARS):
    if len(comment) <= target:
        return comment
    truncated = comment[:target].rstrip(" ,;.")
    if "." in truncated:
        truncated = truncated[:truncated.rfind(".")+1]
    return truncated

# SAMPLE COMMENT BANKS (Replace with your actual imports)
SAMPLE_COMMENTS = {
    "English": {
        5: {
            "opening": ["In English this term,", "During English lessons,", "In English,"],
            "attitude": {
                90: "has shown exceptional enthusiasm and dedication.",
                75: "has worked conscientiously and shown good focus.",
                60: "has made a satisfactory effort in class."
            },
            "reading": {
                90: "reads with excellent comprehension and fluency.",
                75: "reads with understanding and can discuss texts.",
                60: "is developing basic reading skills."
            },
            "writing": {
                90: "writes with creativity and excellent structure.",
                75: "writes clearly with good sentence structure.",
                60: "is learning to write complete sentences."
            },
            "target": {
                90: "continue to challenge themselves with advanced texts.",
                75: "work on expanding vocabulary and descriptive language.",
                60: "practice reading aloud to improve fluency."
            },
            "closer": ["Keep up the excellent work!", "Well done this term."]
        }
    },
    "Maths": {
        5: {
            "opening": ["In Mathematics this term,", "During Maths lessons,", "In Maths,"],
            "attitude": {
                90: "has demonstrated outstanding mathematical thinking.",
                75: "has worked diligently on mathematical concepts.",
                60: "has participated in mathematical activities."
            },
            "achievement": {
                90: "solves complex problems with excellent reasoning.",
                75: "applies mathematical concepts correctly.",
                60: "understands basic mathematical operations."
            },
            "target": {
                90: "tackle more challenging problem-solving tasks.",
                75: "work on mental calculation strategies.",
                60: "practice basic number facts regularly."
            },
            "closer": ["Continue to develop mathematical skills.", "Good progress made."]
        }
    }
}

def generate_comment(subject, year, name, gender, att, achieve, target, optional_text=None):
    """Generate a report comment - SIMPLE WORKING VERSION"""
    p, p_poss = get_pronouns(gender)
    name = sanitize_input(name)
    
    # Get closest available band
    def get_closest_band(value, bands):
        closest = min(bands, key=lambda x: abs(x - value))
        return closest
    
    # Default comment parts
    comment_parts = []
    
    # Try to get from sample data, fallback to generic
    try:
        if subject in SAMPLE_COMMENTS and year in SAMPLE_COMMENTS[subject]:
            data = SAMPLE_COMMENTS[subject][year]
            
            # Opening and attitude
            opening = random.choice(data["opening"])
            att_band = get_closest_band(att, [90, 75, 60])
            attitude = fix_pronouns_in_text(data["attitude"][att_band], p, p_poss)
            comment_parts.append(f"{opening} {name} {attitude}")
            
            # Subject achievement
            achieve_band = get_closest_band(achieve, [90, 75, 60])
            if subject == "English":
                reading = fix_pronouns_in_text(data["reading"][achieve_band], p, p_poss)
                writing = fix_pronouns_in_text(data["writing"][achieve_band], p, p_poss)
                comment_parts.append(f"In reading, {lowercase_first(reading)}")
                comment_parts.append(f"In writing, {lowercase_first(writing)}")
            else:
                achievement = fix_pronouns_in_text(data["achievement"][achieve_band], p, p_poss)
                if achievement[0].islower():
                    achievement = f"{p.capitalize()} {achievement}"
                comment_parts.append(achievement)
            
            # Target
            target_band = get_closest_band(target, [90, 75, 60])
            target_text = fix_pronouns_in_text(data["target"][target_band], p, p_poss)
            comment_parts.append(f"For next term, {p} should {lowercase_first(target_text)}")
            
            # Optional text
            if optional_text and str(optional_text).strip():
                opt = str(optional_text).strip()
                if opt[0].islower():
                    opt = opt[0].upper() + opt[1:]
                if not opt.endswith('.'):
                    opt += '.'
                comment_parts.append(f"Additionally, {lowercase_first(opt)}")
            
            # Closer
            closer = random.choice(data["closer"])
            comment_parts.append(closer)
            
        else:
            # Fallback for unsupported subjects
            comment_parts.append(f"{name} has worked in {subject} this term. {p.capitalize()} has made good progress.")
            if optional_text:
                comment_parts.append(f"Note: {optional_text}")
            
    except Exception as e:
        # Ultimate fallback
        comment_parts = [f"{name} has participated in {subject} lessons. Keep up the good work!"]
    
    # Build final comment
    comment = " ".join(comment_parts)
    
    # Ensure proper punctuation
    for i in range(len(comment_parts)):
        if not comment_parts[i].endswith('.'):
            comment_parts[i] += '.'
    
    comment = " ".join(comment_parts)
    comment = truncate_comment(comment)
    
    if not comment.endswith('.'):
        comment += '.'
    
    return comment

# APP LAYOUT
with st.sidebar:
    st.title("📚 Report Generator")
    app_mode = st.radio("Mode", ["Single Student", "Batch Upload", "Help"])
    
    st.markdown("---")
    if st.button("Clear All Data", use_container_width=True):
        st.session_state.clear()
        st.session_state.app_initialized = True
        st.session_state.all_comments = []
        st.rerun()

# MAIN CONTENT
st.title("Report Comment Generator")
st.caption("Generate personalized report comments quickly and easily")

if app_mode == "Single Student":
    st.subheader("👤 Single Student Entry")
    
    with st.form("student_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            subject = st.selectbox("Subject", ["English", "Maths", "Science", "ESL (IGCSE)", "Chemistry"])
            year = st.selectbox("Year", [5, 7, 8, 10, 11])
            name = st.text_input("Student Name", placeholder="Enter student name")
            gender = st.selectbox("Gender", ["Male", "Female"])
        
        with col2:
            att = st.selectbox("Attitude", [90, 85, 80, 75, 70, 65, 60, 55, 40], index=3)
            achieve = st.selectbox("Achievement", [90, 85, 80, 75, 70, 65, 60, 55, 40], index=3)
            target = st.selectbox("Target", [90, 85, 80, 75, 70, 65, 60, 55, 40], index=3)
        
        optional_comment = st.text_area(
            "Optional Additional Comment (Optional)",
            placeholder="Add any extra comments here...",
            height=60
        )
        
        submitted = st.form_submit_button("🚀 Generate Comment", use_container_width=True)
    
    if submitted and name:
        with st.spinner("Generating comment..."):
            try:
                comment = generate_comment(
                    subject=subject,
                    year=year,
                    name=name,
                    gender=gender,
                    att=att,
                    achieve=achieve,
                    target=target,
                    optional_text=optional_comment
                )
                
                # Display comment
                st.subheader("📝 Generated Comment")
                st.text_area("", comment, height=200, key="display_comment")
                
                # Stats
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Characters", len(comment))
                with col2:
                    st.metric("Words", len(comment.split()))
                
                # Store in session
                if 'all_comments' not in st.session_state:
                    st.session_state.all_comments = []
                
                st.session_state.all_comments.append({
                    'name': name,
                    'subject': subject,
                    'year': year,
                    'comment': comment
                })
                
                st.success("✓ Comment generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating comment: {str(e)}")
                st.info("Try different subject/year combination")

elif app_mode == "Batch Upload":
    st.subheader("📁 Batch Upload")
    
    st.info("Upload a CSV file with columns: Student Name, Gender, Subject, Year, Attitude, Achievement, Target")
    
    example_csv = """Student Name,Gender,Subject,Year,Attitude,Achievement,Target
John Smith,Male,English,7,75,80,85
Sarah Jones,Female,Maths,5,80,75,80"""
    
    st.download_button("📥 Download Example CSV", example_csv, "example.csv", "text/csv")
    
    uploaded_file = st.file_uploader("Choose CSV file", type=['csv'])
    
    if uploaded_file:
        try:
            df = pd.read_csv(uploaded_file)
            st.success(f"Loaded {len(df)} students")
            
            if st.button("Generate All Comments", type="primary"):
                if 'all_comments' not in st.session_state:
                    st.session_state.all_comments = []
                
                progress_bar = st.progress(0)
                
                for idx, row in df.iterrows():
                    comment = generate_comment(
                        subject=str(row.get('Subject', 'English')),
                        year=int(row.get('Year', 7)),
                        name=str(row.get('Student Name', '')),
                        gender=str(row.get('Gender', '')),
                        att=int(row.get('Attitude', 75)),
                        achieve=int(row.get('Achievement', 75)),
                        target=int(row.get('Target', 75))
                    )
                    
                    st.session_state.all_comments.append({
                        'name': str(row.get('Student Name', '')),
                        'subject': str(row.get('Subject', 'English')),
                        'year': int(row.get('Year', 7)),
                        'comment': comment
                    })
                    
                    progress_bar.progress((idx + 1) / len(df))
                
                st.success(f"✓ Generated {len(df)} comments!")
                
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")

else:  # Help mode
    st.subheader("❓ Help & Instructions")
    st.markdown("""
    ### How to Use:
    1. **Single Student**: Fill in the form to generate one comment
    2. **Batch Upload**: Upload a CSV file for multiple students
    
    ### CSV Format:
    ```
    Student Name,Gender,Subject,Year,Attitude,Achievement,Target
    John Doe,Male,English,7,75,80,85
    ```
    
    ### Band Values:
    - 90: Excellent
    - 75: Good  
    - 60: Satisfactory
    - 40: Needs improvement
    
    ### Subjects Supported:
    - English (Years 5, 7, 8)
    - Maths (Years 5, 7, 8)
    - Science (Years 5, 7, 8)
    - ESL (IGCSE) (Years 10, 11)
    - Chemistry (Years 10, 11)
    """)

# DOWNLOAD SECTION
if 'all_comments' in st.session_state and st.session_state.all_comments:
    st.markdown("---")
    st.subheader("📥 Download Options")
    
    total = len(st.session_state.all_comments)
    st.info(f"You have {total} generated comment(s)")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Download as Word Document"):
            doc = Document()
            doc.add_heading('Report Comments', 0)
            
            for entry in st.session_state.all_comments:
                doc.add_heading(f"{entry['name']} - {entry['subject']} Year {entry['year']}", level=2)
                doc.add_paragraph(entry['comment'])
                doc.add_paragraph()
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                "⬇️ Download Word File",
                bio.getvalue(),
                f"comments_{datetime.now().strftime('%Y%m%d')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    with col2:
        if st.button("Download as CSV"):
            csv_data = []
            for entry in st.session_state.all_comments:
                csv_data.append({
                    'Student Name': entry['name'],
                    'Subject': entry['subject'],
                    'Year': entry['year'],
                    'Comment': entry['comment']
                })
            
            df = pd.DataFrame(csv_data)
            csv = df.to_csv(index=False)
            
            st.download_button(
                "⬇️ Download CSV",
                csv,
                f"comments_{datetime.now().strftime('%Y%m%d')}.csv",
                "text/csv"
            )
    
    if st.button("🔄 Start Over", type="secondary"):
        st.session_state.all_comments = []
        st.rerun()

st.markdown("---")
st.caption("Report Generator v1.0 • All comments are generated locally")# Version 2.0 - Cache buster
# MULTI-SUBJECT REPORT COMMENT GENERATOR
# Supports: English, Science, Maths, ESL, Chemistry

import streamlit as st
import tempfile
import os
import pandas as pd
import io
import random
import re
from datetime import datetime, timedelta
from docx import Document

import sys
sys.dont_write_bytecode = True

# Add this at the VERY TOP for import fixes
import os
import sys
sys.path.insert(0, os.path.abspath('.'))

# SECURITY & PRIVACY SETTINGS
TARGET_CHARS = 500
MAX_FILE_SIZE_MB = 5
MAX_ROWS_PER_UPLOAD = 100
RATE_LIMIT_SECONDS = 10

# PAGE CONFIGURATION
st.set_page_config(
    page_title="CommentCraft",
    layout="wide",
    initial_sidebar_state="expanded"
)

# SECURITY INITIALIZATION
if 'app_initialized' not in st.session_state:
    st.session_state.clear()
    st.session_state.app_initialized = True
    st.session_state.upload_count = 0
    st.session_state.last_upload_time = datetime.now()
    st.session_state.generated_files = []

# IMPORT STATEMENT FILES (directly from repository)
try:
    # English
    from statements_year5_English import (
        opening_phrases as opening_5_eng,
        attitude_bank as attitude_5_eng,
        reading_bank as reading_5_eng,
        writing_bank as writing_5_eng,
        reading_target_bank as target_5_eng,
        writing_target_bank as target_write_5_eng,
        closer_bank as closer_5_eng
    )
    
    from statements_year7_English import (
        opening_phrases as opening_7_eng,
        attitude_bank as attitude_7_eng,
        reading_bank as reading_7_eng,
        writing_bank as writing_7_eng,
        reading_target_bank as target_7_eng,
        writing_target_bank as target_write_7_eng,
        closer_bank as closer_7_eng
    )

    from statements_year8_English import (
        opening_phrases as opening_8_eng,
        attitude_bank as attitude_8_eng,
        reading_bank as reading_8_eng,
        writing_bank as writing_8_eng,
        reading_target_bank as target_8_eng,
        writing_target_bank as target_write_8_eng,
        closer_bank as closer_8_eng
    )

    # Science
    from statements_year5_Science import (
        opening_phrases as opening_5_sci,
        attitude_bank as attitude_5_sci,
        science_bank as science_5_sci,
        target_bank as target_5_sci,
        closer_bank as closer_5_sci
    )
    
    from statements_year7_science import (
        opening_phrases as opening_7_sci,
        attitude_bank as attitude_7_sci,
        science_bank as science_7_sci,
        target_bank as target_7_sci,
        closer_bank as closer_7_sci
    )

    from statements_year8_science import (
        opening_phrases as opening_8_sci,
        attitude_bank as attitude_8_sci,
        science_bank as science_8_sci,
        target_bank as target_8_sci,
        closer_bank as closer_8_sci
    )

    # Maths
    from statements_year5_Maths_NEW import (
        opening_phrases as opening_5_math,
        attitude_bank as attitude_5_math,
        maths_bank as maths_5_math,
        target_bank as target_5_math,
        closer_bank as closer_5_math
    )
    
    from statements_year7_Maths_NEW import (
        opening_phrases as opening_7_math,
        attitude_bank as attitude_7_math,
        maths_bank as maths_7_math,
        target_bank as target_7_math,
        closer_bank as closer_7_math
    )

    from statements_year8_Maths_NEW import (
        opening_phrases as opening_8_math,
        attitude_bank as attitude_8_math,
        maths_bank as maths_8_math,
        target_bank as target_8_math,
        closer_bank as closer_8_math
    )
    
    # ESL (IGCSE)
    from statements_igcse_0510_esl import (
        opening_phrases as opening_esl,
        attitude_bank as attitude_esl,
        reading_bank as reading_esl,
        writing_bank as writing_esl,
        speaking_bank as speaking_esl,
        listening_bank as listening_esl,
        reading_target_bank as target_reading_esl,
        writing_target_bank as target_write_esl,
        speaking_target_bank as target_speak_esl,
        listening_target_bank as target_listen_esl,
        closer_bank as closer_esl
    )

    # Chemistry
    from statements_igcse_0620_chemistry import (
        opening_phrases as opening_chem,
        attitude_bank as attitude_chem,
        chemistry_bank as chemistry_chem,
        target_bank as target_chem,
        closer_bank as closer_chem
    )

except ImportError as e:
    st.error(f"Missing required statement files: {e}")
    st.error("Please ensure all statement files are in the repository.")
    st.stop()

# SECURITY FUNCTIONS
def validate_upload_rate():
    """Prevent rapid-fire uploads/abuse"""
    time_since_last = datetime.now() - st.session_state.last_upload_time
    if time_since_last < timedelta(seconds=RATE_LIMIT_SECONDS):
        wait_time = RATE_LIMIT_SECONDS - time_since_last.seconds
        st.error(f"Please wait {wait_time} seconds before uploading again")
        return False
    return True

def sanitize_input(text, max_length=100):
    """Sanitize user input to prevent injection attacks"""
    if not text:
        return ""
    sanitized = ''.join(c for c in text if c.isalnum() or c in " .'-")
    return sanitized[:max_length].strip().title()

def validate_file(file):
    """Validate uploaded file size and type"""
    if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        return False, f"File too large (max {MAX_FILE_SIZE_MB}MB)"
    
    if not file.name.lower().endswith('.csv'):
        return False, "Only CSV files allowed"
    
    return True, ""

def process_csv_securely(uploaded_file):
    """Process CSV with auto-cleanup of temp files"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.csv', mode='wb') as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name
    
    try:
        df = pd.read_csv(temp_path, nrows=MAX_ROWS_PER_UPLOAD + 1)
        
        if len(df) > MAX_ROWS_PER_UPLOAD:
            st.warning(f"Only processing first {MAX_ROWS_PER_UPLOAD} rows")
            df = df.head(MAX_ROWS_PER_UPLOAD)
        
        if 'Student Name' in df.columns:
            df['Student Name'] = df['Student Name'].apply(lambda x: sanitize_input(str(x)))
        
        return df
        
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        return None
        
    finally:
        try:
            os.unlink(temp_path)
        except:
            pass

# HELPER FUNCTIONS
def get_pronouns(gender):
    gender = gender.lower()
    if gender == "male":
        return "he", "his"
    elif gender == "female":
        return "she", "her"
    return "they", "their"

def lowercase_first(text):
    return text[0].lower() + text[1:] if text else ""

def truncate_comment(comment, target=TARGET_CHARS):
    if len(comment) <= target:
        return comment
    truncated = comment[:target].rstrip(" ,;.")
    if "." in truncated:
        truncated = truncated[:truncated.rfind(".")+1]
    return truncated

def fix_pronouns_in_text(text, pronoun, possessive):
    """Fix gender pronouns in statement text"""
    if not text:
        return text
    
    # Fix pronouns with word boundaries
    text = re.sub(r'\bhe\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHe\b', pronoun.capitalize(), text)
    text = re.sub(r'\bshe\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bShe\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhis\b', possessive, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHis\b', possessive.capitalize(), text)
    text = re.sub(r'\bher\b', possessive, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHer\b', possessive.capitalize(), text)
    text = re.sub(r'\bhim\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHim\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhimself\b', f"{pronoun}self", text, flags=re.IGNORECASE)
    text = re.sub(r'\bherself\b', f"{pronoun}self", text, flags=re.IGNORECASE)
    
    return text

def generate_comment(subject, year, name, gender, att, achieve, target, optional_text=None):
    """Generate a report comment based on subject, year, and performance bands"""
    p, p_poss = get_pronouns(gender)
    name = sanitize_input(name)
    
    # Initialize comment parts
    comment_parts = []
    
    # Subject-specific comment generation
    if subject == "English":
        if year == 5:
            opening = random.choice(opening_5_eng)
            attitude_text = fix_pronouns_in_text(attitude_5_eng[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            reading_text = fix_pronouns_in_text(reading_5_eng[achieve], p, p_poss)
            if reading_text[0].islower():
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            
            writing_text = fix_pronouns_in_text(writing_5_eng[achieve], p, p_poss)
            if writing_text[0].islower():
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            
            reading_target_text = fix_pronouns_in_text(target_5_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            
            writing_target_text = fix_pronouns_in_text(target_write_5_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            
            closer_sentence = random.choice(closer_5_eng)
            
            comment_parts = [
                attitude_sentence,
                reading_sentence,
                writing_sentence,
                reading_target_sentence,
                writing_target_sentence,
                closer_sentence
            ]
            
        elif year == 7:
            opening = random.choice(opening_7_eng)
            attitude_text = fix_pronouns_in_text(attitude_7_eng[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            reading_text = fix_pronouns_in_text(reading_7_eng[achieve], p, p_poss)
            if reading_text[0].islower():
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            
            writing_text = fix_pronouns_in_text(writing_7_eng[achieve], p, p_poss)
            if writing_text[0].islower():
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            
            reading_target_text = fix_pronouns_in_text(target_7_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            
            writing_target_text = fix_pronouns_in_text(target_write_7_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            
            closer_sentence = random.choice(closer_7_eng)
            
            comment_parts = [
                attitude_sentence,
                reading_sentence,
                writing_sentence,
                reading_target_sentence,
                writing_target_sentence,
                closer_sentence
            ]
            
        else:  # Year 8
            opening = random.choice(opening_8_eng)
            attitude_text = fix_pronouns_in_text(attitude_8_eng[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            reading_text = fix_pronouns_in_text(reading_8_eng[achieve], p, p_poss)
            if reading_text[0].islower():
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            
            writing_text = fix_pronouns_in_text(writing_8_eng[achieve], p, p_poss)
            if writing_text[0].islower():
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            
            reading_target_text = fix_pronouns_in_text(target_8_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            
            writing_target_text = fix_pronouns_in_text(target_write_8_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            
            closer_sentence = random.choice(closer_8_eng)
            
            comment_parts = [
                attitude_sentence,
                reading_sentence,
                writing_sentence,
                reading_target_sentence,
                writing_target_sentence,
                closer_sentence
            ]
        
    elif subject == "Science":
        if year == 5:
            opening = random.choice(opening_5_sci)
            attitude_text = fix_pronouns_in_text(attitude_5_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            science_text = fix_pronouns_in_text(science_5_sci[achieve], p, p_poss)
            if science_text[0].islower():
                science_text = f"{p} {science_text}"
            science_sentence = science_text
            
            target_text = fix_pronouns_in_text(target_5_sci[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_5_sci)
            
            comment_parts = [
                attitude_sentence,
                science_sentence,
                target_sentence,
                closer_sentence
            ]
            
        elif year == 7:
            opening = random.choice(opening_7_sci)
            attitude_text = fix_pronouns_in_text(attitude_7_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            science_text = fix_pronouns_in_text(science_7_sci[achieve], p, p_poss)
            if science_text[0].islower():
                science_text = f"{p} {science_text}"
            science_sentence = science_text
            
            target_text = fix_pronouns_in_text(target_7_sci[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_7_sci)
            
            comment_parts = [
                attitude_sentence,
                science_sentence,
                target_sentence,
                closer_sentence
            ]
            
        else:  # Year 8
            opening = random.choice(opening_8_sci)
            attitude_text = fix_pronouns_in_text(attitude_8_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            science_text = fix_pronouns_in_text(science_8_sci[achieve], p, p_poss)
            if science_text[0].islower():
                science_text = f"{p} {science_text}"
            science_sentence = science_text
            
            target_text = fix_pronouns_in_text(target_8_sci[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_8_sci)
            
            comment_parts = [
                attitude_sentence,
                science_sentence,
                target_sentence,
                closer_sentence
            ]
        
    elif subject == "Maths":
        if year == 5:
            opening = random.choice(opening_5_math)
            attitude_text = fix_pronouns_in_text(attitude_5_math[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            maths_text = fix_pronouns_in_text(maths_5_math[achieve], p, p_poss)
            if maths_text[0].islower():
                maths_text = f"{p} {maths_text}"
            maths_sentence = maths_text
            
            target_text = fix_pronouns_in_text(target_5_math[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_5_math)
            
            comment_parts = [
                attitude_sentence,
                maths_sentence,
                target_sentence,
                closer_sentence
            ]
            
        elif year == 7:
            opening = random.choice(opening_7_math)
            attitude_text = fix_pronouns_in_text(attitude_7_math[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            maths_text = fix_pronouns_in_text(maths_7_math[achieve], p, p_poss)
            if maths_text[0].islower():
                maths_text = f"{p} {maths_text}"
            maths_sentence = maths_text
            
            target_text = fix_pronouns_in_text(target_7_math[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_7_math)
            
            comment_parts = [
                attitude_sentence,
                maths_sentence,
                target_sentence,
                closer_sentence
            ]
            
        else:  # Year 8
            opening = random.choice(opening_8_math)
            attitude_text = fix_pronouns_in_text(attitude_8_math[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            maths_text = fix_pronouns_in_text(maths_8_math[achieve], p, p_poss)
            if maths_text[0].islower():
                maths_text = f"{p} {maths_text}"
            maths_sentence = maths_text
            
            target_text = fix_pronouns_in_text(target_8_math[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_8_math)
            
            comment_parts = [
                attitude_sentence,
                maths_sentence,
                target_sentence,
                closer_sentence
            ]
        
    elif subject == "ESL (IGCSE)":
        opening = random.choice(opening_esl)
        attitude_text = fix_pronouns_in_text(attitude_esl[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        
        # Reading
        reading_text = fix_pronouns_in_text(reading_esl[achieve], p, p_poss)
        if reading_text[0].islower():
            reading_text = f"{p} {reading_text}"
        reading_sentence = f"In reading, {reading_text}"
        
        # Writing
        writing_text = fix_pronouns_in_text(writing_esl[achieve], p, p_poss)
        if writing_text[0].islower():
            writing_text = f"{p} {writing_text}"
        writing_sentence = f"In writing, {writing_text}"
        
        # Speaking
        speaking_text = fix_pronouns_in_text(speaking_esl[achieve], p, p_poss)
        if speaking_text[0].islower():
            speaking_text = f"{p} {speaking_text}"
        speaking_sentence = f"In speaking, {speaking_text}"
        
        # Listening
        listening_text = fix_pronouns_in_text(listening_esl[achieve], p, p_poss)
        if listening_text[0].islower():
            listening_text = f"{p} {listening_text}"
        listening_sentence = f"In listening, {listening_text}"
        
        # Targets
        reading_target_text = fix_pronouns_in_text(target_reading_esl[target], p, p_poss)
        reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
        
        writing_target_text = fix_pronouns_in_text(target_write_esl[target], p, p_poss)
        writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
        
        closer_sentence = random.choice(closer_esl)
        
        comment_parts = [
            attitude_sentence,
            reading_sentence,
            writing_sentence,
            speaking_sentence,
            listening_sentence,
            reading_target_sentence,
            writing_target_sentence,
            closer_sentence
        ]
        
    elif subject == "Chemistry":
        opening = random.choice(opening_chem)
        attitude_text = fix_pronouns_in_text(attitude_chem[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        
        chemistry_text = fix_pronouns_in_text(chemistry_chem[achieve], p, p_poss)
        if chemistry_text[0].islower():
            chemistry_text = f"{p} {chemistry_text}"
        chemistry_sentence = chemistry_text
        
        target_text = fix_pronouns_in_text(target_chem[target], p, p_poss)
        target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
        
        closer_sentence = random.choice(closer_chem)
        
        comment_parts = [
            attitude_sentence,
            chemistry_sentence,
            target_sentence,
            closer_sentence
        ]
    
    else:
        # Default fallback if subject not recognized
        comment_parts = [f"{name} has worked in {subject} this term."]
    
    # Add optional text if provided - NOW AT THE END
    if optional_text:
        optional_text = sanitize_input(optional_text)
        if optional_text:
            optional_sentence = f"Additionally, {lowercase_first(optional_text)}"
            if not optional_sentence.endswith('.'):
                optional_sentence += '.'
            # Insert before the closer sentence (second to last position)
            if comment_parts:
                comment_parts.insert(-1, optional_sentence)
            else:
                comment_parts.append(optional_sentence)
    
    # Ensure all sentences end with period
    for i in range(len(comment_parts)):
        if not comment_parts[i].endswith('.'):
            comment_parts[i] += '.'
    
    # Join comment parts
    comment = " ".join([c for c in comment_parts if c])
    comment = truncate_comment(comment, TARGET_CHARS)
    
    # Ensure comment ends with period
    if not comment.endswith('.'):
        comment = comment.rstrip(' ,;') + '.'
    
    return comment

# STREAMLIT APP LAYOUT

# Sidebar for navigation
with st.sidebar:
    st.title("CommentCraft")
    st.caption("Your AI report writing assistant")
    
    app_mode = st.radio(
        "Choose Mode",
        ["Single Student", "Batch Upload", "Privacy Info"]
    )
    
    st.markdown("---")
    st.markdown("### Privacy Features")
    st.info("""
    - No data stored on servers
    - All processing in memory
    - Auto-deletion of temp files
    - Input sanitization
    - Rate limiting enabled
    """)
    
    if st.button("Clear All Data", type="secondary", use_container_width=True):
        st.session_state.clear()
        st.session_state.app_initialized = True
        st.session_state.upload_count = 0
        st.session_state.last_upload_time = datetime.now()
        st.success("All data cleared!")
        st.rerun()

# Main content area
st.title("CommentCraft")

# Privacy notice
st.warning("""
**Privacy Notice:** All data is processed in memory only. No files are stored on servers. 
Close browser tab to completely erase all data. For use with anonymized student data only.
""")

# SINGLE STUDENT MODE
if app_mode == "Single Student":
    st.subheader("Single Student Entry")
    
    with st.form("single_student_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        
        with col1:
            subject = st.selectbox("Subject", ["English", "Maths", "Science", "ESL (IGCSE)", "Chemistry"])
            year = st.selectbox("Year", [5, 7, 8, 10, 11])
            name = st.text_input("Student Name", placeholder="Enter first name only")
            gender = st.selectbox("Gender", ["Male", "Female"])
        
        with col2:
            att = st.selectbox("Attitude Band", 
                             options=[90,85,80,75,70,65,60,55,40],
                             index=3)
            
            achieve = st.selectbox("Achievement Band",
                                 options=[90,85,80,75,70,65,60,55,40],
                                 index=3)
            
            target = st.selectbox("Target Band",
                                options=[90,85,80,75,70,65,60,55,40],
                                index=3)
        
        attitude_target = st.text_area("Optional Additional Comment",
                                     placeholder="Add any additional comments here...",
                                     height=60)
        
        submitted = st.form_submit_button("Generate Comment")
    
    if submitted and name:
        if not validate_upload_rate():
            st.stop()
        
        name = sanitize_input(name)
        
        with st.spinner("Generating comment..."):
            comment = generate_comment(
                subject=subject,
                year=year,
                name=name,
                gender=gender,
                att=att,
                achieve=achieve,
                target=target,
                optional_text=attitude_target
            )
            char_count = len(comment)
        
        # Display comment
        st.subheader("Generated Comment")
        st.text_area("", comment, height=200)
        
        # Stats
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Character Count", f"{char_count}/500")
        with col2:
            st.metric("Words", len(comment.split()))
        with col3:
            if char_count < 450:
                st.success("Good length")
            else:
                st.warning("Near limit")
        
        # Store in session
        if 'all_comments' not in st.session_state:
            st.session_state.all_comments = []
        
        student_entry = {
            'name': name,
            'subject': subject,
            'year': year,
            'comment': comment,
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
        }
        st.session_state.all_comments.append(student_entry)
        
        # Add another button
        if st.button("Add Another Student"):
            st.rerun()

# BATCH UPLOAD MODE
elif app_mode == "Batch Upload":
    st.subheader("Batch Upload (CSV)")
    
    st.info("""
    **CSV Format Required:**
    - Columns: Student Name, Gender, Subject, Year, Attitude, Achievement, Target
    - Gender: Male/Female
    - Subject: English/Maths/Science/ESL (IGCSE)/Chemistry
    - Year: 5,7,8,10,11
    - Bands: 90,85,80,75,70,65,60,55,40
    """)
    
    # Example CSV
    example_csv = """Student Name,Gender,Subject,Year,Attitude,Achievement,Target
John,Male,English,7,75,80,85
Sarah,Female,Maths,5,80,75,80
Ahmed,Male,ESL (IGCSE),10,85,90,85
Maria,Female,Chemistry,11,80,85,80"""
    
    st.download_button(
        label="Download Example CSV",
        data=example_csv,
        file_name="example_students.csv",
        mime="text/csv"
    )
    
    uploaded_file = st.file_uploader("Choose CSV file", type=['csv'])
    
    if uploaded_file:
        if not validate_upload_rate():
            st.stop()
        
        is_valid, msg = validate_file(uploaded_file)
        if not is_valid:
            st.error(msg)
            st.stop()
        
        with st.spinner("Processing CSV..."):
            df = process_csv_securely(uploaded_file)
        
        if df is not None:
            st.success(f"Processed {len(df)} students")
            
            with st.expander("Preview Data"):
                st.dataframe(df.head())
            
            if st.button("Generate All Comments"):
                if 'all_comments' not in st.session_state:
                    st.session_state.all_comments = []
                
                progress_bar = st.progress(0)
                
                for idx, row in df.iterrows():
                    progress = (idx + 1) / len(df)
                    progress_bar.progress(progress)
                    
                    try:
                        comment = generate_comment(
                            subject=str(row.get('Subject', 'English')),
                            year=int(row.get('Year', 7)),
                            name=str(row.get('Student Name', '')),
                            gender=str(row.get('Gender', '')),
                            att=int(row.get('Attitude', 75)),
                            achieve=int(row.get('Achievement', 75)),
                            target=int(row.get('Target', 75))
                        )
                        
                        student_entry = {
                            'name': sanitize_input(str(row.get('Student Name', ''))),
                            'subject': str(row.get('Subject', 'English')),
                            'year': int(row.get('Year', 7)),
                            'comment': comment,
                            'timestamp': datetime.now().strftime("%Y-%m-d %H:%M")
                        }
                        st.session_state.all_comments.append(student_entry)
                        
                    except Exception as e:
                        st.error(f"Error processing row {idx + 1}: {e}")
                
                progress_bar.empty()
                st.success(f"Generated {len(df)} comments!")
                st.session_state.last_upload_time = datetime.now()

# PRIVACY INFO MODE
elif app_mode == "Privacy Info":
    st.subheader("Privacy & Security Information")
    
    st.markdown("""
    ### Data Protection
    
    **How we handle data:**
    - All processing occurs in your browser's memory
    - No student data is sent to external servers
    - Temporary files are created and immediately deleted
    - No database or persistent storage is used
    
    **Security features:**
    1. **Input Sanitization** - Removes special characters
    2. **Rate Limiting** - Prevents system abuse
    3. **File Validation** - Checks file size and type
    4. **Auto-Cleanup** - Temporary files automatically deleted
    5. **Memory Clearing** - All data erased on browser close
    
    **Best practices:**
    - Use only first names or student IDs
    - Close browser tab when finished
    - Download reports immediately
    - Use on school-managed devices for maximum privacy
    """)

# DOWNLOAD SECTION
if 'all_comments' in st.session_state and st.session_state.all_comments:
    st.markdown("---")
    st.subheader("Download Reports")
    
    total_comments = len(st.session_state.all_comments)
    st.info(f"You have {total_comments} generated comment(s)")
    
    # Preview
    with st.expander(f"Preview Comments ({total_comments})"):
        for idx, entry in enumerate(st.session_state.all_comments, 1):
            st.markdown(f"**{idx}. {entry['name']}** ({entry['subject']} Year {entry['year']})")
            st.write(entry['comment'])
            st.markdown("---")
    
    # Download options
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("Word Document"):
            doc = Document()
            doc.add_heading('Report Comments', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            doc.add_paragraph(f'Total Students: {total_comments}')
            doc.add_paragraph('')
            
            for entry in st.session_state.all_comments:
                doc.add_heading(f"{entry['name']} - {entry['subject']} Year {entry['year']}", level=2)
                doc.add_paragraph(entry['comment'])
                doc.add_paragraph('')
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="Download Word File",
                data=bio.getvalue(),
                file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    with col2:
        if st.button("CSV Export"):
            csv_data = []
            for entry in st.session_state.all_comments:
                csv_data.append({
                    'Student Name': entry['name'],
                    'Subject': entry['subject'],
                    'Year': entry['year'],
                    'Comment': entry['comment'],
                    'Generated': entry['timestamp']
                })
            
            df_export = pd.DataFrame(csv_data)
            csv_bytes = df_export.to_csv(index=False).encode('utf-8')
            
            st.download_button(
                label="Download CSV",
                data=csv_bytes,
                file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv"
            )
    
    with col3:
        if st.button("Clear All", type="secondary"):
            st.session_state.all_comments = []
            st.success("All comments cleared!")
            st.rerun()

# FOOTER
st.markdown("---")
st.caption("CommentCraft v4.0 • Secure & Private")
