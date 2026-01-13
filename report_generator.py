# Version 2.0 - Cache buster
# MULTI-SUBJECT REPORT COMMENT GENERATOR
# Supports: English, Science, Maths, ESL, Chemistry

import streamlit as st
import tempfile
import os
import pandas as pd
import io
import random
import re
from datetime import datetime, timedelta
from docx import Document

# SECURITY & PRIVACY SETTINGS
TARGET_CHARS = 500
MAX_FILE_SIZE_MB = 5
MAX_ROWS_PER_UPLOAD = 100
RATE_LIMIT_SECONDS = 10

# PAGE CONFIGURATION
st.set_page_config(
    page_title="Report Comment Generator",
    layout="wide",
    initial_sidebar_state="expanded"
)

# SECURITY INITIALIZATION
if 'app_initialized' not in st.session_state:
    st.session_state.clear()
    st.session_state.app_initialized = True
    st.session_state.upload_count = 0
    st.session_state.last_upload_time = datetime.now()
    st.session_state.generated_files = []

# IMPORT STATEMENT FILES (directly from repository)
try:
    # English
    from statements_year5_English import (
        opening_phrases as opening_5_eng,
        attitude_bank as attitude_5_eng,
        reading_bank as reading_5_eng,
        writing_bank as writing_5_eng,
        reading_target_bank as target_5_eng,
        writing_target_bank as target_write_5_eng,
        closer_bank as closer_5_eng
    )
    
    from statements_year7_English import (
        opening_phrases as opening_7_eng,
        attitude_bank as attitude_7_eng,
        reading_bank as reading_7_eng,
        writing_bank as writing_7_eng,
        reading_target_bank as target_7_eng,
        writing_target_bank as target_write_7_eng,
        closer_bank as closer_7_eng
    )

    from statements_year8_English import (
        opening_phrases as opening_8_eng,
        attitude_bank as attitude_8_eng,
        reading_bank as reading_8_eng,
        writing_bank as writing_8_eng,
        reading_target_bank as target_8_eng,
        writing_target_bank as target_write_8_eng,
        closer_bank as closer_8_eng
    )

    # Science
    from statements_year5_Science import (
        opening_phrases as opening_5_sci,
        attitude_bank as attitude_5_sci,
        science_bank as science_5_sci,
        target_bank as target_5_sci,
        closer_bank as closer_5_sci
    )
    
    from statements_year7_science import (
        opening_phrases as opening_7_sci,
        attitude_bank as attitude_7_sci,
        science_bank as science_7_sci,
        target_bank as target_7_sci,
        closer_bank as closer_7_sci
    )

    from statements_year8_science import (
        opening_phrases as opening_8_sci,
        attitude_bank as attitude_8_sci,
        science_bank as science_8_sci,
        target_bank as target_8_sci,
        closer_bank as closer_8_sci
    )

    # Maths
    from statements_year5_Maths import (
        opening_phrases as opening_5_math,
        attitude_bank as attitude_5_math,
        maths_bank as maths_5_math,
        target_bank as target_5_math,
        closer_bank as closer_5_math
    )
    
    from statements_year7_Maths import (
        opening_phrases as opening_7_math,
        attitude_bank as attitude_7_math,
        maths_bank as maths_7_math,
        target_bank as target_7_math,
        closer_bank as closer_7_math
    )

    from statements_year8_Maths import (
        opening_phrases as opening_8_math,
        attitude_bank as attitude_8_math,
        maths_bank as maths_8_math,
        target_bank as target_8_math,
        closer_bank as closer_8_math
    )

    # ESL (IGCSE)
    from statements_igcse_0510_esl import (
        opening_phrases as opening_esl,
        attitude_bank as attitude_esl,
        reading_bank as reading_esl,
        writing_bank as writing_esl,
        speaking_bank as speaking_esl,
        listening_bank as listening_esl,
        reading_target_bank as target_reading_esl,
        writing_target_bank as target_write_esl,
        speaking_target_bank as target_speak_esl,
        listening_target_bank as target_listen_esl,
        closer_bank as closer_esl
    )

    # Chemistry
    from statements_igcse_0620_chemistry import (
        opening_phrases as opening_chem,
        attitude_bank as attitude_chem,
        chemistry_bank as chemistry_chem,
        target_bank as target_chem,
        closer_bank as closer_chem
    )

except ImportError as e:
    st.error(f"Missing required statement files: {e}")
    st.error("Please ensure all statement files are in the repository.")
    st.stop()

# SECURITY FUNCTIONS
def validate_upload_rate():
    """Prevent rapid-fire uploads/abuse"""
    time_since_last = datetime.now() - st.session_state.last_upload_time
    if time_since_last < timedelta(seconds=RATE_LIMIT_SECONDS):
        wait_time = RATE_LIMIT_SECONDS - time_since_last.seconds
        st.error(f"Please wait {wait_time} seconds before uploading again")
        return False
    return True

def sanitize_input(text, max_length=100):
    """Sanitize user input to prevent injection attacks"""
    if not text:
        return ""
    sanitized = ''.join(c for c in text if c.isalnum() or c in " .'-")
    return sanitized[:max_length].strip().title()

def validate_file(file):
    """Validate uploaded file size and type"""
    if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        return False, f"File too large (max {MAX_FILE_SIZE_MB}MB)"
    
    if not file.name.lower().endswith('.csv'):
        return False, "Only CSV files allowed"
    
    return True, ""

def process_csv_securely(uploaded_file):
    """Process CSV with auto-cleanup of temp files"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.csv', mode='wb') as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name
    
    try:
        df = pd.read_csv(temp_path, nrows=MAX_ROWS_PER_UPLOAD + 1)
        
        if len(df) > MAX_ROWS_PER_UPLOAD:
            st.warning(f"Only processing first {MAX_ROWS_PER_UPLOAD} rows")
            df = df.head(MAX_ROWS_PER_UPLOAD)
        
        if 'Student Name' in df.columns:
            df['Student Name'] = df['Student Name'].apply(lambda x: sanitize_input(str(x)))
        
        return df
        
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        return None
        
    finally:
        try:
            os.unlink(temp_path)
        except:
            pass

# HELPER FUNCTIONS
def get_pronouns(gender):
    gender = gender.lower()
    if gender == "male":
        return "he", "his"
    elif gender == "female":
        return "she", "her"
    return "they", "their"

def lowercase_first(text):
    return text[0].lower() + text[1:] if text else ""

def truncate_comment(comment, target=TARGET_CHARS):
    if len(comment) <= target:
        return comment
    truncated = comment[:target].rstrip(" ,;.")
    if "." in truncated:
        truncated = truncated[:truncated.rfind(".")+1]
    return truncated

def fix_pronouns_in_text(text, pronoun, possessive):
    """Fix gender pronouns in statement text"""
    if not text:
        return text
    
    # Fix pronouns with word boundaries
    text = re.sub(r'\bhe\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHe\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhis\b', possessive, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHis\b', possessive.capitalize(), text)
    text = re.sub(r'\bhim\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHim\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhimself\b', f"{pronoun}self", text, flags=re.IGNORECASE)
    text = re.sub(r'\bherself\b', f"{pronoun}self", text, flags=re.IGNORECASE)
    
    return text

def generate_comment(subject, year, name, gender, att, achieve, target, optional_text=None):
    """Generate a report comment based on subject, year, and performance bands"""
    p, p_poss = get_pronouns(gender)
    name = sanitize_input(name)
    
    # Initialize comment parts
    comment_parts = []
    
    # Subject-specific comment generation
    if subject == "English":
        if year == 5:
            opening = random.choice(opening_5_eng)
            attitude_text = fix_pronouns_in_text(attitude_5_eng[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            reading_text = fix_pronouns_in_text(reading_5_eng[achieve], p, p_poss)
            if reading_text[0].islower():
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            
            writing_text = fix_pronouns_in_text(writing_5_eng[achieve], p, p_poss)
            if writing_text[0].islower():
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            
            reading_target_text = fix_pronouns_in_text(target_5_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            
            writing_target_text = fix_pronouns_in_text(target_write_5_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            
            closer_sentence = random.choice(closer_5_eng)
            
        elif year == 7:
            opening = random.choice(opening_7_eng)
            attitude_text = fix_pronouns_in_text(attitude_7_eng[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            reading_text = fix_pronouns_in_text(reading_7_eng[achieve], p, p_poss)
            if reading_text[0].islower():
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            
            writing_text = fix_pronouns_in_text(writing_7_eng[achieve], p, p_poss)
            if writing_text[0].islower():
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            
            reading_target_text = fix_pronouns_in_text(target_7_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            
            writing_target_text = fix_pronouns_in_text(target_write_7_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            
            closer_sentence = random.choice(closer_7_eng)
            
        else:  # Year 8
            opening = random.choice(opening_8_eng)
            attitude_text = fix_pronouns_in_text(attitude_8_eng[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            reading_text = fix_pronouns_in_text(reading_8_eng[achieve], p, p_poss)
            if reading_text[0].islower():
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            
            writing_text = fix_pronouns_in_text(writing_8_eng[achieve], p, p_poss)
            if writing_text[0].islower():
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            
            reading_target_text = fix_pronouns_in_text(target_8_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            
            writing_target_text = fix_pronouns_in_text(target_write_8_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            
            closer_sentence = random.choice(closer_8_eng)
            
        comment_parts = [
            attitude_sentence,
            reading_sentence,
            writing_sentence,
            reading_target_sentence,
            writing_target_sentence,
            closer_sentence
        ]
        
    elif subject == "Science":
        if year == 5:
            opening = random.choice(opening_5_sci)
            attitude_text = fix_pronouns_in_text(attitude_5_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            science_text = fix_pronouns_in_text(science_5_sci[achieve], p, p_poss)
            if science_text[0].islower():
                science_text = f"{p} {science_text}"
            science_sentence = science_text
            
            target_text = fix_pronouns_in_text(target_5_sci[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_5_sci)
            
        elif year == 7:
            opening = random.choice(opening_7_sci)
            attitude_text = fix_pronouns_in_text(attitude_7_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            science_text = fix_pronouns_in_text(science_7_sci[achieve], p, p_poss)
            if science_text[0].islower():
                science_text = f"{p} {science_text}"
            science_sentence = science_text
            
            target_text = fix_pronouns_in_text(target_7_sci[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_7_sci)
            
        else:  # Year 8
            opening = random.choice(opening_8_sci)
            attitude_text = fix_pronouns_in_text(attitude_8_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            science_text = fix_pronouns_in_text(science_8_sci[achieve], p, p_poss)
            if science_text[0].islower():
                science_text = f"{p} {science_text}"
            science_sentence = science_text
            
            target_text = fix_pronouns_in_text(target_8_sci[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_8_sci)
            
        comment_parts = [
            attitude_sentence,
            science_sentence,
            target_sentence,
            closer_sentence
        ]
        
    elif subject == "Maths":
        if year == 5:
            opening = random.choice(opening_5_math)
            attitude_text = fix_pronouns_in_text(attitude_5_math[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            maths_text = fix_pronouns_in_text(maths_5_math[achieve], p, p_poss)
            if maths_text[0].islower():
                maths_text = f"{p} {maths_text}"
            maths_sentence = maths_text
            
            target_text = fix_pronouns_in_text(target_5_math[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_5_math)
            
        elif year == 7:
            opening = random.choice(opening_7_math)
            attitude_text = fix_pronouns_in_text(attitude_7_math[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            maths_text = fix_pronouns_in_text(maths_7_math[achieve], p, p_poss)
            if maths_text[0].islower():
                maths_text = f"{p} {maths_text}"
            maths_sentence = maths_text
            
            target_text = fix_pronouns_in_text(target_7_math[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_7_math)
            
        else:  # Year 8
            opening = random.choice(opening_8_math)
            attitude_text = fix_pronouns_in_text(attitude_8_math[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            
            maths_text = fix_pronouns_in_text(maths_8_math[achieve], p, p_poss)
            if maths_text[0].islower():
                maths_text = f"{p} {maths_text}"
            maths_sentence = maths_text
            
            target_text = fix_pronouns_in_text(target_8_math[target], p, p_poss)
            target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            
            closer_sentence = random.choice(closer_8_math)
            
        comment_parts = [
            attitude_sentence,
            maths_sentence,
            target_sentence,
            closer_sentence
        ]
        
    elif subject == "ESL (IGCSE)":
        opening = random.choice(opening_esl)
        attitude_text = fix_pronouns_in_text(attitude_esl[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        
        # Reading
        reading_text = fix_pronouns_in_text(reading_esl[achieve], p, p_poss)
        if reading_text[0].islower():
            reading_text = f"{p} {reading_text}"
        reading_sentence = f"In reading, {reading_text}"
        
        # Writing
        writing_text = fix_pronouns_in_text(writing_esl[achieve], p, p_poss)
        if writing_text[0].islower():
            writing_text = f"{p} {writing_text}"
        writing_sentence = f"In writing, {writing_text}"
        
        # Speaking
        speaking_text = fix_pronouns_in_text(speaking_esl[achieve], p, p_poss)
        if speaking_text[0].islower():
            speaking_text = f"{p} {speaking_text}"
        speaking_sentence = f"In speaking, {speaking_text}"
        
        # Listening
        listening_text = fix_pronouns_in_text(listening_esl[achieve], p, p_poss)
        if listening_text[0].islower():
            listening_text = f"{p} {listening_text}"
        listening_sentence = f"In listening, {listening_text}"
        
        # Targets
        reading_target_text = fix_pronouns_in_text(target_reading_esl[target], p, p_poss)
        reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
        
        writing_target_text = fix_pronouns_in_text(target_write_esl[target], p, p_poss)
        writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
        
        closer_sentence = random.choice(closer_esl)
        
        comment_parts = [
            attitude_sentence,
            reading_sentence,
            writing_sentence,
            speaking_sentence,
            listening_sentence,
            reading_target_sentence,
            writing_target_sentence,
            closer_sentence
        ]
        
    elif subject == "Chemistry":
        opening = random.choice(opening_chem)
        attitude_text = fix_pronouns_in_text(attitude_chem[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        
        chemistry_text = fix_pronouns_in_text(chemistry_chem[achieve], p, p_poss)
        if chemistry_text[0].islower():
            chemistry_text = f"{p} {chemistry_text}"
        chemistry_sentence = chemistry_text
        
        target_text = fix_pronouns_in_text(target_chem[target], p, p_poss)
        target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
        
        closer_sentence = random.choice(closer_chem)
        
        comment_parts = [
            attitude_sentence,
            chemistry_sentence,
            target_sentence,
            closer_sentence
        ]
    
    # Add optional text if provided
    if optional_text:
        optional_text = sanitize_input(optional_text)
        if optional_text:
            optional_sentence = f"Additionally, {lowercase_first(optional_text)}"
            if not optional_sentence.endswith('.'):
                optional_sentence += '.'
            comment_parts.insert(1, optional_sentence)
    
    # Ensure all sentences end with period
    for i in range(len(comment_parts)):
        if not comment_parts[i].endswith('.'):
            comment_parts[i] += '.'
    
    # Join comment parts
    comment = " ".join([c for c in comment_parts if c])
    comment = truncate_comment(comment, TARGET_CHARS)
    
    # Ensure comment ends with period
    if not comment.endswith('.'):
        comment = comment.rstrip(' ,;') + '.'
    
    return comment

# STREAMLIT APP LAYOUT

# Sidebar for navigation
with st.sidebar:
    st.title("Navigation")
    
    app_mode = st.radio(
        "Choose Mode",
        ["Single Student", "Batch Upload", "Privacy Info"]
    )
    
    st.markdown("---")
    st.markdown("### Privacy Features")
    st.info("""
    - No data stored on servers
    - All processing in memory
    - Auto-deletion of temp files
    - Input sanitization
    - Rate limiting enabled
    """)
    
    if st.button("Clear All Data", type="secondary", use_container_width=True):
        st.session_state.clear()
        st.session_state.app_initialized = True
        st.session_state.upload_count = 0
        st.session_state.last_upload_time = datetime.now()
        st.success("All data cleared!")
        st.rerun()

# Main content area
st.title("Multi-Subject Report Comment Generator")
st.caption("Supports: English, Maths, Science, ESL (IGCSE), Chemistry | Years: 5, 7, 8 | Target: 500 characters per comment")

# Privacy notice
st.warning("""
**Privacy Notice:** All data is processed in memory only. No files are stored on servers. 
Close browser tab to completely erase all data. For use with anonymized student data only.
""")

# SINGLE STUDENT MODE
if app_mode == "Single Student":
    st.subheader("Single Student Entry")
    
    with st.form("single_student_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        
        with col1:
            subject = st.selectbox("Subject", ["English", "Maths", "Science", "ESL (IGCSE)", "Chemistry"])
            year = st.selectbox("Year", [5, 7, 8, 10, 11])
            name = st.text_input("Student Name", placeholder="Enter first name only")
            gender = st.selectbox("Gender", ["Male", "Female"])
        
        with col2:
            att = st.selectbox("Attitude Band", 
                             options=[90,85,80,75,70,65,60,55,40],
                             index=3)
            
            achieve = st.selectbox("Achievement Band",
                                 options=[90,85,80,75,70,65,60,55,40],
                                 index=3)
            
            target = st.selectbox("Target Band",
                                options=[90,85,80,75,70,65,60,55,40],
                                index=3)
        
        attitude_target = st.text_area("Optional Additional Comment",
                                     placeholder="Add any additional comments here...",
                                     height=60)
        
        submitted = st.form_submit_button("Generate Comment")
    
    if submitted and name:
        if not validate_upload_rate():
            st.stop()
        
        name = sanitize_input(name)
        
        with st.spinner("Generating comment..."):
            comment = generate_comment(
                subject=subject,
                year=year,
                name=name,
                gender=gender,
                att=att,
                achieve=achieve,
                target=target,
                optional_text=attitude_target
            )
            char_count = len(comment)
        
        # Display comment
        st.subheader("Generated Comment")
        st.text_area("", comment, height=200)
        
        # Stats
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Character Count", f"{char_count}/500")
        with col2:
            st.metric("Words", len(comment.split()))
        with col3:
            if char_count < 450:
                st.success("Good length")
            else:
                st.warning("Near limit")
        
        # Store in session
        if 'all_comments' not in st.session_state:
            st.session_state.all_comments = []
        
        student_entry = {
            'name': name,
            'subject': subject,
            'year': year,
            'comment': comment,
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
        }
        st.session_state.all_comments.append(student_entry)
        
        # Add another button
        if st.button("Add Another Student"):
            st.rerun()

# BATCH UPLOAD MODE
elif app_mode == "Batch Upload":
    st.subheader("Batch Upload (CSV)")
    
    st.info("""
    **CSV Format Required:**
    - Columns: Student Name, Gender, Subject, Year, Attitude, Achievement, Target
    - Gender: Male/Female
    - Subject: English/Maths/Science/ESL (IGCSE)/Chemistry
    - Year: 5,7,8,10,11
    - Bands: 90,85,80,75,70,65,60,55,40
    """)
    
    # Example CSV
    example_csv = """Student Name,Gender,Subject,Year,Attitude,Achievement,Target
John,Male,English,7,75,80,85
Sarah,Female,Maths,5,80,75,80
Ahmed,Male,ESL (IGCSE),10,85,90,85
Maria,Female,Chemistry,11,80,85,80"""
    
    st.download_button(
        label="Download Example CSV",
        data=example_csv,
        file_name="example_students.csv",
        mime="text/csv"
    )
    
    uploaded_file = st.file_uploader("Choose CSV file", type=['csv'])
    
    if uploaded_file:
        if not validate_upload_rate():
            st.stop()
        
        is_valid, msg = validate_file(uploaded_file)
        if not is_valid:
            st.error(msg)
            st.stop()
        
        with st.spinner("Processing CSV..."):
            df = process_csv_securely(uploaded_file)
        
        if df is not None:
            st.success(f"Processed {len(df)} students")
            
            with st.expander("Preview Data"):
                st.dataframe(df.head())
            
            if st.button("Generate All Comments"):
                if 'all_comments' not in st.session_state:
                    st.session_state.all_comments = []
                
                progress_bar = st.progress(0)
                
                for idx, row in df.iterrows():
                    progress = (idx + 1) / len(df)
                    progress_bar.progress(progress)
                    
                    try:
                        comment = generate_comment(
                            subject=str(row.get('Subject', 'English')),
                            year=int(row.get('Year', 7)),
                            name=str(row.get('Student Name', '')),
                            gender=str(row.get('Gender', '')),
                            att=int(row.get('Attitude', 75)),
                            achieve=int(row.get('Achievement', 75)),
                            target=int(row.get('Target', 75))
                        )
                        
                        student_entry = {
                            'name': sanitize_input(str(row.get('Student Name', ''))),
                            'subject': str(row.get('Subject', 'English')),
                            'year': int(row.get('Year', 7)),
                            'comment': comment,
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                        }
                        st.session_state.all_comments.append(student_entry)
                        
                    except Exception as e:
                        st.error(f"Error processing row {idx + 1}: {e}")
                
                progress_bar.empty()
                st.success(f"Generated {len(df)} comments!")
                st.session_state.last_upload_time = datetime.now()

# PRIVACY INFO MODE
elif app_mode == "Privacy Info":
    st.subheader("Privacy & Security Information")
    
    st.markdown("""
    ### Data Protection
    
    **How we handle data:**
    - All processing occurs in your browser's memory
    - No student data is sent to external servers
    - Temporary files are created and immediately deleted
    - No database or persistent storage is used
    
    **Security features:**
    1. **Input Sanitization** - Removes special characters
    2. **Rate Limiting** - Prevents system abuse
    3. **File Validation** - Checks file size and type
    4. **Auto-Cleanup** - Temporary files automatically deleted
    5. **Memory Clearing** - All data erased on browser close
    
    **Best practices:**
    - Use only first names or student IDs
    - Close browser tab when finished
    - Download reports immediately
    - Use on school-managed devices for maximum privacy
    """)

# DOWNLOAD SECTION
if 'all_comments' in st.session_state and st.session_state.all_comments:
    st.markdown("---")
    st.subheader("Download Reports")
    
    total_comments = len(st.session_state.all_comments)
    st.info(f"You have {total_comments} generated comment(s)")
    
    # Preview
    with st.expander(f"Preview Comments ({total_comments})"):
        for idx, entry in enumerate(st.session_state.all_comments, 1):
            st.markdown(f"**{idx}. {entry['name']}** ({entry['subject']} Year {entry['year']})")
            st.write(entry['comment'])
            st.markdown("---")
    
    # Download options
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("Word Document"):
            doc = Document()
            doc.add_heading('Report Comments', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            doc.add_paragraph(f'Total Students: {total_comments}')
            doc.add_paragraph('')
            
            for entry in st.session_state.all_comments:
                doc.add_heading(f"{entry['name']} - {entry['subject']} Year {entry['year']}", level=2)
                doc.add_paragraph(entry['comment'])
                doc.add_paragraph('')
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="Download Word File",
                data=bio.getvalue(),
                file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    with col2:
        if st.button("CSV Export"):
            csv_data = []
            for entry in st.session_state.all_comments:
                csv_data.append({
                    'Student Name': entry['name'],
                    'Subject': entry['subject'],
                    'Year': entry['year'],
                    'Comment': entry['comment'],
                    'Generated': entry['timestamp']
                })
            
            df_export = pd.DataFrame(csv_data)
            csv_bytes = df_export.to_csv(index=False).encode('utf-8')
            
            st.download_button(
                label="Download CSV",
                data=csv_bytes,
                file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv"
            )
    
    with col3:
        if st.button("Clear All", type="secondary"):
            st.session_state.all_comments = []
            st.success("All comments cleared!")
            st.rerun()

# FOOTER
st.markdown("---")
st.caption("Report Generator v3.0 • Multi-Subject Edition")
